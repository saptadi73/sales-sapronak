"""
Script untuk membuat dokumentasi fungsional aplikasi Vue Odoo Sales dalam format Word (.docx)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_background(cell, hex_color):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    if level == 1:
        run.font.color.rgb = RGBColor(0x10, 0x5A, 0x3B)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x1A, 0x7A, 0x52)
    else:
        run.font.color.rgb = RGBColor(0x2A, 0x9A, 0x6A)
    return heading

def add_info_box(doc, text, bg_color="E8F5E9"):
    """Add a shaded info box paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    # Add shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg_color)
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_code_block(doc, code_text):
    """Add a code-style paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F5F9')
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table_with_header(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, '105A3B')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'F0FAF5'
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_background(cell, bg)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

def build_document():
    doc = Document()

    # --- Page margins ---
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ===== COVER PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('DOKUMENTASI FUNGSIONAL')
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0x10, 0x5A, 0x3B)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run('Aplikasi Frontend Vue.js – Odoo Sales')
    sub_run.bold = True
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x1A, 0x7A, 0x52)

    doc.add_paragraph()

    app_para = doc.add_paragraph()
    app_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    app_run = app_para.add_run('vue-odoo-sales')
    app_run.font.size = Pt(13)
    app_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    app_run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.add_run(f'Tanggal Dokumen: {datetime.date.today().strftime("%d %B %Y")}')

    doc.add_page_break()

    # ===== DAFTAR ISI =====
    add_heading(doc, 'Daftar Isi', level=1)
    toc_items = [
        '1.  Pendahuluan',
        '2.  Tujuan Aplikasi',
        '3.  Arsitektur & Teknologi',
        '4.  Alur Autentikasi (Login)',
        '5.  Fitur-Fitur Utama',
        '    5.1  Pembuatan QR Customer (QR Generator)',
        '    5.2  Pembacaan QR Customer (QR Reader)',
        '    5.3  Pembuatan Sales Order (Sales Order View)',
        '6.  Navigasi & Routing',
        '7.  State Management (Session Store)',
        '8.  Ringkasan Endpoint API',
        '9.  Tipe Data Utama',
        '10. Penanganan Error & Notifikasi',
        '11. Konfigurasi & Setup Awal',
        '12. Catatan Keamanan',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5 if item.startswith('    ') else 0)
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ===== 1. PENDAHULUAN =====
    add_heading(doc, '1. Pendahuluan', level=1)
    doc.add_paragraph(
        'Dokumen ini merupakan dokumentasi fungsional untuk aplikasi frontend Vue.js yang '
        'berfungsi sebagai antarmuka pengguna (UI) untuk modul penjualan berbasis Odoo. '
        'Aplikasi ini dibangun menggunakan Vue 3, TypeScript, Vite, dan Tailwind CSS, '
        'serta berkomunikasi dengan backend Odoo melalui JSON-RPC via REST API kustom '
        'yang tersedia pada modul Odoo grt_sales_business_category.'
    )
    doc.add_paragraph(
        'Aplikasi ini dirancang untuk digunakan oleh tim sales di lapangan, memungkinkan '
        'akses cepat ke data customer melalui QR code, pengecekan posisi akuntansi customer, '
        'riwayat transaksi, dan pembuatan Sales Order langsung dari perangkat mobile maupun desktop.'
    )

    # ===== 2. TUJUAN APLIKASI =====
    add_heading(doc, '2. Tujuan Aplikasi', level=1)
    doc.add_paragraph('Aplikasi ini memiliki tujuan fungsional sebagai berikut:')
    goals = [
        'Menyediakan antarmuka login yang terintegrasi dengan sistem autentikasi session Odoo.',
        'Menghasilkan QR Code customer yang dapat dicetak atau disimpan dalam format PNG.',
        'Membaca QR Code customer menggunakan kamera perangkat untuk mengidentifikasi customer.',
        'Menampilkan informasi detail customer, ringkasan hutang/piutang, dan riwayat Sales Order.',
        'Memfasilitasi pembuatan draft Sales Order dengan berbagai jenis bon (bon kering, bon partus, bon reguler, non-ongkir).',
        'Mendukung multi-item per Sales Order dengan fitur pencarian produk.',
        'Mendukung kalkulasi grand total otomatis sebelum order dikirim ke Odoo.',
    ]
    for g in goals:
        p = doc.add_paragraph(g, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # ===== 3. ARSITEKTUR & TEKNOLOGI =====
    add_heading(doc, '3. Arsitektur & Teknologi', level=1)
    doc.add_paragraph(
        'Aplikasi menggunakan arsitektur Single Page Application (SPA) berbasis komponen Vue 3 '
        'dengan Composition API. Semua komunikasi ke backend dilakukan secara asinkron melalui '
        'Fetch API dengan format JSON.'
    )
    doc.add_paragraph()

    add_heading(doc, '3.1  Stack Teknologi', level=2)
    tech_rows = [
        ['Vue 3 + Composition API', 'Framework UI utama berbasis komponen'],
        ['TypeScript', 'Pengetikan statis untuk keamanan tipe data'],
        ['Vite', 'Build tool dan development server'],
        ['Tailwind CSS v4', 'Utility-first CSS framework'],
        ['Pinia', 'State management (session store)'],
        ['Vue Router v4', 'Client-side routing dengan navigation guard'],
        ['qrcode', 'Library generate QR Code ke data URL (PNG)'],
        ['@zxing/browser', 'Library scan QR Code dari kamera perangkat'],
        ['Fetch API', 'HTTP client native browser untuk komunikasi ke Odoo'],
    ]
    add_table_with_header(doc, ['Teknologi', 'Fungsi'], tech_rows, col_widths=[5, 10])

    doc.add_paragraph()
    add_heading(doc, '3.2  Struktur Direktori Utama', level=2)
    add_code_block(doc,
        'src/\n'
        '  views/        # Halaman utama (Login, QR Generator, QR Reader, Sales Order)\n'
        '  components/   # Komponen reusable (Navbar, QrScanner, ProductSearchInput)\n'
        '  services/     # Lapisan komunikasi API (odooApi.ts)\n'
        '  stores/       # State global Pinia (session.ts)\n'
        '  types/        # TypeScript interface & type (odoo.ts)\n'
        '  utils/        # Utilitas umum (qr.ts: extract QR ref, format currency)\n'
        '  router/       # Konfigurasi routing Vue Router\n'
        '  assets/css/   # Global stylesheet'
    )

    doc.add_page_break()

    # ===== 4. ALUR AUTENTIKASI =====
    add_heading(doc, '4. Alur Autentikasi (Login)', level=1)
    doc.add_paragraph(
        'Autentikasi menggunakan mekanisme session Odoo melalui endpoint JSON-RPC. '
        'Tidak ada JWT atau token bearer — aplikasi mengandalkan cookie session yang dikelola '
        'browser secara otomatis.'
    )

    add_heading(doc, '4.1  Halaman Login', level=2)
    doc.add_paragraph('Pengguna mengisi tiga field pada form login:')
    login_fields = [
        ('Base URL Odoo', 'URL server Odoo, contoh: https://your-odoo-host'),
        ('Database', 'Nama database Odoo yang digunakan'),
        ('Username / Email', 'Email atau login akun Odoo'),
        ('Password', 'Password akun Odoo'),
    ]
    add_table_with_header(doc, ['Field', 'Keterangan'], login_fields, col_widths=[5, 10])

    doc.add_paragraph()
    add_heading(doc, '4.2  Alur Login', level=2)
    login_steps = [
        'Pengguna mengisi form dan menekan tombol Login.',
        'Frontend memanggil POST /api/sales/authenticate dengan credentials.',
        'Odoo mengembalikan session_id, uid, nama user, company, dan partner_id.',
        'Data session disimpan ke Pinia store (sessionStore) dan dipersist di localStorage.',
        'Browser menerima cookie session otomatis dari Odoo.',
        'Pengguna diarahkan ke halaman QR Generator (default) atau halaman tujuan sebelumnya.',
        'Jika login gagal, pesan error ditampilkan pada form.',
    ]
    for i, step in enumerate(login_steps, 1):
        p = doc.add_paragraph(f'{i}. {step}')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    add_heading(doc, '4.3  Navigation Guard', level=2)
    doc.add_paragraph(
        'Vue Router dikonfigurasi dengan navigation guard yang memeriksa status autentikasi '
        'sebelum setiap navigasi. Halaman dengan meta requiresAuth: true hanya dapat diakses '
        'jika session valid tersedia. Jika session tidak ada, pengguna diarahkan ke halaman '
        'login dengan parameter query redirect sehingga setelah login berhasil, pengguna '
        'dikembalikan ke halaman yang semula dituju.'
    )

    add_info_box(doc,
        'Catatan: Seluruh request ke endpoint yang membutuhkan autentikasi menggunakan '
        'credentials: "include" agar cookie session dikirim ke Odoo. Konfigurasi CORS '
        'dan cookie policy perlu diaktifkan di sisi server Odoo jika frontend berjalan '
        'di domain berbeda.'
    )

    doc.add_page_break()

    # ===== 5. FITUR-FITUR UTAMA =====
    add_heading(doc, '5. Fitur-Fitur Utama', level=1)

    # 5.1 QR Generator
    add_heading(doc, '5.1  Pembuatan QR Customer (QR Generator)', level=2)
    doc.add_paragraph(
        'Fitur ini memungkinkan pengguna membuat QR Code unik untuk customer Odoo. '
        'QR Code yang dihasilkan mengandung referensi customer (customer_qr_ref) yang '
        'dapat dipindai untuk mengidentifikasi customer secara instan.'
    )

    doc.add_paragraph('Mode Pencarian Customer:')
    modes = [
        ('Customer ID', 'Pengguna memasukkan ID numerik customer Odoo (partner_id).'),
        ('SR Reference', 'Pengguna memasukkan kode referensi customer (customer_qr_ref).'),
    ]
    add_table_with_header(doc, ['Mode', 'Keterangan'], modes, col_widths=[4, 11])

    doc.add_paragraph()
    doc.add_paragraph('Format Konten QR:')
    formats = [
        ('ref', 'QR berisi string referensi saja, contoh: CUSTQR2603-000001. Lebih ringkas dan mudah dipindai.'),
        ('json', 'QR berisi payload JSON dengan customer_id, customer_qr_ref, dan customer_name. Cocok jika scanner ingin membaca metadata langsung.'),
    ]
    add_table_with_header(doc, ['Format', 'Keterangan'], formats, col_widths=[2, 13])

    doc.add_paragraph()
    doc.add_paragraph('Setelah QR berhasil dibuat, pengguna dapat:')
    qr_actions = [
        'Melihat preview QR Code di halaman.',
        'Mengunduh gambar QR dalam format PNG.',
        'Mencetak QR melalui dialog print browser (nama customer dan referensi tercetak di atas gambar QR).',
    ]
    for a in qr_actions:
        p = doc.add_paragraph(a, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # 5.2 QR Reader
    add_heading(doc, '5.2  Pembacaan QR Customer (QR Reader)', level=2)
    doc.add_paragraph(
        'Fitur ini memungkinkan pengguna memindai QR Code customer menggunakan kamera '
        'perangkat, kemudian secara otomatis mengambil informasi lengkap customer dari Odoo.'
    )

    doc.add_paragraph('Cara penggunaan:')
    reader_steps = [
        'Pengguna membuka halaman QR Reader.',
        'Komponen QrScanner mengaktifkan kamera perangkat (meminta izin kamera).',
        'Scanner menggunakan library @zxing/browser untuk mendeteksi QR Code secara real-time.',
        'Saat QR terdeteksi, nilai referensi customer diekstrak dari konten QR.',
        'Frontend secara paralel memanggil tiga endpoint: detail customer, ringkasan akuntansi, dan riwayat Sales Order.',
        'Informasi ditampilkan dalam tiga kartu terpisah di halaman.',
        'Pengguna dapat memilih untuk langsung membuat Sales Order untuk customer tersebut.',
    ]
    for i, step in enumerate(reader_steps, 1):
        p = doc.add_paragraph(f'{i}. {step}')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    doc.add_paragraph('Input Manual:')
    doc.add_paragraph(
        'Selain scan kamera, pengguna juga dapat memasukkan nilai QR secara manual pada '
        'textarea (mendukung format ref string maupun JSON payload).'
    )

    doc.add_paragraph()
    doc.add_paragraph('Informasi yang ditampilkan setelah scan:')
    info_sections = [
        ('Detail Customer', 'Nama, alamat, nomor telepon, email, wilayah pengiriman, dan jangka waktu pembayaran.'),
        ('Ringkasan Akuntansi', 'Total piutang (receivable), total hutang (payable), dan tabel aging hutang/piutang dalam rentang 1-10, 11-30, 31-60, 61-90 hari, dan >90 hari.'),
        ('Riwayat Sales Order', 'Daftar 20 Sales Order terakhir customer: nomor order, tanggal, tanggal komitmen, total, dan status order.'),
    ]
    add_table_with_header(doc, ['Bagian', 'Konten'], info_sections, col_widths=[4, 11])

    doc.add_paragraph()

    # 5.3 Sales Order View
    add_heading(doc, '5.3  Pembuatan Sales Order (Sales Order View)', level=2)
    doc.add_paragraph(
        'Halaman ini memungkinkan pengguna membuat draft Sales Order di Odoo langsung '
        'dari frontend. Customer diidentifikasi melalui QR Code (dapat dari scan maupun '
        'query parameter dari halaman QR Reader).'
    )

    add_heading(doc, 'Jenis Bon (Tipe Sales Order)', level=3)
    bon_types = [
        ('Bon Kering', 'bon-kering', 'Sales Order dengan auto ongkos kirim. Wajib memiliki wilayah pengiriman.'),
        ('Bon Partus', 'bon-partus', 'Sales Order dengan auto ongkos kirim. Wajib memiliki wilayah pengiriman.'),
        ('Bon Reguler', 'bon-reguler', 'Sales Order dengan auto ongkos kirim. Wajib memiliki wilayah pengiriman.'),
        ('Non Ongkir', 'non-ongkir', 'Sales Order tanpa ongkos kirim otomatis.'),
    ]
    add_table_with_header(doc,
        ['Label', 'Kode API', 'Keterangan'],
        bon_types,
        col_widths=[3.5, 3.5, 8]
    )

    doc.add_paragraph()
    add_heading(doc, 'Alur Pembuatan Sales Order', level=3)
    so_steps = [
        'Halaman terbuka dengan customer yang sudah terisi jika ada query parameter ?qr=...',
        'Pengguna memilih jenis bon dari dropdown.',
        'Pengguna mengisi tanggal komitmen (commitment_date) opsional.',
        'Pengguna memilih payment term dari daftar yang diambil dari Odoo.',
        'Pengguna menambahkan item produk menggunakan komponen ProductSearchInput (autocomplete).',
        'Setiap baris memiliki field: produk, jumlah (qty), dan harga satuan (price_unit).',
        'Grand total dihitung otomatis dari semua baris.',
        'Pengguna menekan tombol Buat Order untuk mengirim ke Odoo.',
        'Jika berhasil, nomor Sales Order dan total ditampilkan.',
    ]
    for i, step in enumerate(so_steps, 1):
        p = doc.add_paragraph(f'{i}. {step}')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    add_heading(doc, 'Validasi', level=3)
    validations = [
        'Customer wajib terisi (dari QR scan atau input manual).',
        'Jenis bon wajib dipilih.',
        'Jika jenis bon membutuhkan ongkir (bukan non-ongkir), customer harus memiliki shipping_wilayah_id.',
        'Minimal satu baris produk harus ada.',
        'Setiap baris produk harus memiliki product_id yang valid.',
    ]
    for v in validations:
        p = doc.add_paragraph(v, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ===== 6. NAVIGASI & ROUTING =====
    add_heading(doc, '6. Navigasi & Routing', level=1)
    doc.add_paragraph(
        'Routing dikelola oleh Vue Router v4 dengan mode history (HTML5 History API). '
        'Redirect otomatis ke /qr-generator saat mengakses root URL.'
    )
    doc.add_paragraph()

    route_rows = [
        ['/login', 'LoginView', 'Tidak', 'Halaman login Odoo'],
        ['/ (root)', '—', 'Tidak', 'Redirect ke /qr-generator'],
        ['/qr-generator', 'QrGeneratorView', 'Ya', 'Buat QR Code customer'],
        ['/qr-reader', 'QrReaderView', 'Ya', 'Scan QR & tampilkan info customer'],
        ['/sales-order', 'SalesOrderView', 'Ya', 'Buat draft Sales Order'],
    ]
    add_table_with_header(doc,
        ['Path', 'Komponen View', 'Requires Auth', 'Fungsi'],
        route_rows,
        col_widths=[3.5, 4, 3, 6]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Navbar (AppNavbar) menampilkan menu navigasi yang hanya terlihat saat pengguna '
        'sudah login. Navbar menampilkan nama user dan nama company yang sedang aktif.'
    )

    # ===== 7. STATE MANAGEMENT =====
    add_heading(doc, '7. State Management (Session Store)', level=1)
    doc.add_paragraph(
        'Pinia digunakan sebagai state management. Store utama adalah sessionStore yang '
        'menyimpan data session Odoo aktif dan dipersist ke localStorage.'
    )
    doc.add_paragraph()

    state_rows = [
        ['baseUrl', 'string', 'URL server Odoo yang digunakan'],
        ['db', 'string', 'Nama database Odoo aktif'],
        ['user', 'OdooAuthData | null', 'Data user yang sedang login (uid, nama, company, dll)'],
        ['isAuthenticated', 'boolean (computed)', 'True jika user tersedia'],
    ]
    add_table_with_header(doc,
        ['State', 'Tipe', 'Keterangan'],
        state_rows,
        col_widths=[4, 4.5, 7]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Session dipulihkan dari localStorage setiap kali navigation guard dijalankan '
        '(restoreSession()). Saat logout, session store dikosongkan dan localStorage dibersihkan.'
    )

    doc.add_page_break()

    # ===== 8. RINGKASAN ENDPOINT API =====
    add_heading(doc, '8. Ringkasan Endpoint API', level=1)
    doc.add_paragraph(
        'Semua endpoint menggunakan metode POST dan format JSON-RPC. '
        'Endpoint dengan Auth = user memerlukan cookie session Odoo yang valid.'
    )
    doc.add_paragraph()

    endpoint_rows = [
        ['/api/sales/authenticate', 'Public', 'Login dan membuat session Odoo'],
        ['/api/sales/products', 'User', 'List produk dan harga (global)'],
        ['/api/sales/payment-terms', 'User', 'List Payment Terms'],
        ['/api/sales/customer-qr-by-id', 'User', 'Ambil customer_qr_ref dari customer_id'],
        ['/api/sales/customer-qr-payload-by-id', 'User', 'Payload QR siap render dari customer_id'],
        ['/api/sales/customer-qr-payload-by-ref', 'User', 'Payload QR siap render dari customer_qr_ref'],
        ['/api/sales/customer-detail-by-qr', 'User', 'Detail customer dari customer_qr_ref'],
        ['/api/sales/customer-accounting-summary-by-qr', 'User', 'Ringkasan aging hutang/piutang customer'],
        ['/api/sales/orders-by-qr', 'User', 'Riwayat Sales Order customer'],
        ['/api/sales/susu-olahan/customers', 'User', 'List customer global untuk frontend Susu Olahan'],
        ['/api/sales/susu-olahan/customer-search', 'User', 'Autocomplete customer (typeahead) Susu Olahan'],
        ['/api/sales/susu-olahan/products', 'User', 'List produk saleable kategori Susu Olahan'],
        ['/api/sales/susu-olahan/shipping-products', 'User', 'List produk ongkos kirim Susu Olahan'],
        ['/api/sales/susu-olahan/draft-order', 'User', 'Buat draft Sales Order Susu Olahan dengan auto ongkir'],
        ['/api/sales/minimarket/grid-products', 'User', 'Produk dalam format grid entry untuk UI minimarket'],
        ['/api/sales/minimarket/draft-order', 'User', 'Buat draft Sales Order dari quantity grid minimarket'],
        ['/api/sales/draft-order', 'User', 'Buat draft Sales Order multi-item (umum)'],
        ['/api/sales/draft-order/bon-kering', 'User', 'Buat draft Sales Order bon kering'],
        ['/api/sales/draft-order/bon-partus', 'User', 'Buat draft Sales Order bon partus'],
        ['/api/sales/draft-order/bon-reguler', 'User', 'Buat draft Sales Order bon reguler'],
        ['/api/sales/draft-order/non-ongkir', 'User', 'Buat draft Sales Order tanpa ongkos kirim'],
    ]
    add_table_with_header(doc,
        ['Endpoint', 'Auth', 'Fungsi'],
        endpoint_rows,
        col_widths=[6, 2, 8]
    )

    doc.add_page_break()

    # ===== 9. TIPE DATA UTAMA =====
    add_heading(doc, '9. Tipe Data Utama', level=1)
    doc.add_paragraph(
        'Semua tipe data didefinisikan dalam src/types/odoo.ts menggunakan TypeScript interface. '
        'Berikut adalah tipe data utama yang digunakan pada aplikasi ini.'
    )

    add_heading(doc, '9.1  OdooAuthData', level=2)
    doc.add_paragraph('Data yang dikembalikan setelah login berhasil:')
    auth_fields = [
        ['uid', 'number', 'ID user Odoo'],
        ['session_id', 'string', 'Session ID yang dibuat Odoo'],
        ['db', 'string', 'Nama database'],
        ['login', 'string', 'Email/login user'],
        ['name', 'string', 'Nama lengkap user'],
        ['partner_id', 'number', 'ID partner terkait user'],
        ['company_id', 'number', 'ID perusahaan aktif'],
        ['company_name', 'string', 'Nama perusahaan aktif'],
    ]
    add_table_with_header(doc, ['Field', 'Tipe', 'Keterangan'], auth_fields, col_widths=[4, 3, 8.5])

    doc.add_paragraph()
    add_heading(doc, '9.2  CustomerDetailData', level=2)
    doc.add_paragraph('Detail informasi customer:')
    customer_fields = [
        ['partner_id', 'number', 'ID partner Odoo'],
        ['name', 'string', 'Nama customer'],
        ['customer_qr_ref', 'string', 'Referensi QR unik customer (contoh: CUSTQR2603-000001)'],
        ['street / street2', 'string?', 'Alamat customer'],
        ['city', 'string?', 'Kota customer'],
        ['phone / mobile', 'string?', 'Nomor telepon'],
        ['email', 'string?', 'Email customer'],
        ['shipping_wilayah_id', 'number?', 'ID wilayah pengiriman'],
        ['shipping_wilayah_name', 'string?', 'Nama wilayah pengiriman'],
        ['payment_term_id', 'number?', 'ID jangka waktu pembayaran'],
        ['payment_term_name', 'string?', 'Nama jangka waktu pembayaran'],
    ]
    add_table_with_header(doc, ['Field', 'Tipe', 'Keterangan'], customer_fields, col_widths=[4, 3, 8.5])

    doc.add_paragraph()
    add_heading(doc, '9.3  DraftOrderPayload', level=2)
    doc.add_paragraph('Payload yang dikirim saat membuat Sales Order:')
    draft_fields = [
        ['partner_id', 'number?', 'ID customer (opsional jika memakai customer_qr_ref)'],
        ['customer_qr_ref', 'string?', 'Referensi QR customer (opsional jika memakai partner_id)'],
        ['commitment_date', 'string?', 'Tanggal komitmen pengiriman (format ISO datetime)'],
        ['sale_order_type', 'string?', 'Tipe Sales Order (kering, basah, dll)'],
        ['payment_term_id', 'number?', 'ID Payment Term yang dipilih'],
        ['order_line', 'DraftOrderLineInput[]', 'Daftar item produk yang dipesan'],
    ]
    add_table_with_header(doc, ['Field', 'Tipe', 'Keterangan'], draft_fields, col_widths=[4, 4, 7.5])

    doc.add_paragraph()
    add_heading(doc, '9.4  DraftOrderResult', level=2)
    doc.add_paragraph('Hasil yang dikembalikan setelah Sales Order berhasil dibuat:')
    result_fields = [
        ['sale_order_id', 'number', 'ID Sales Order yang dibuat di Odoo'],
        ['name', 'string', 'Nomor Sales Order (contoh: S/00123)'],
        ['state', 'string', 'Status order (draft, sale, dll)'],
        ['amount_total', 'number', 'Total nilai order'],
        ['line_count', 'number', 'Jumlah baris produk'],
        ['wilayah_id / wilayah_name', 'number? / string?', 'Wilayah pengiriman yang diterapkan'],
        ['shipping_product_id / name', 'number? / string?', 'Produk ongkos kirim yang ditambahkan otomatis'],
        ['shipping_price_per_kg', 'number?', 'Tarif ongkos kirim per kg'],
    ]
    add_table_with_header(doc, ['Field', 'Tipe', 'Keterangan'], result_fields, col_widths=[4.5, 4, 7])

    doc.add_page_break()

    # ===== 10. PENANGANAN ERROR =====
    add_heading(doc, '10. Penanganan Error & Notifikasi', level=1)
    doc.add_paragraph(
        'Aplikasi memiliki sistem penanganan error terpusat melalui custom event ODOO_API_TOAST_EVENT '
        'yang di-dispatch oleh service layer dan di-listen oleh komponen App.vue.'
    )
    doc.add_paragraph()

    error_rows = [
        ['network-error', 'Koneksi ke server Odoo bermasalah — periksa internet atau base URL.'],
        ['session-expired', 'Session Odoo berakhir — pengguna perlu login ulang.'],
        ['http-error (4xx/5xx)', 'Pesan error dari response Odoo ditampilkan langsung di UI.'],
    ]
    add_table_with_header(doc,
        ['Jenis Error', 'Pesan / Aksi'],
        error_rows,
        col_widths=[5, 10.5]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Selain toast notifikasi, setiap halaman juga menampilkan errorMessage inline di bawah '
        'form atau komponen terkait. Pesan sukses (successMessage) juga ditampilkan inline '
        'setelah operasi berhasil.'
    )

    # ===== 11. KONFIGURASI & SETUP AWAL =====
    add_heading(doc, '11. Konfigurasi & Setup Awal', level=1)

    add_heading(doc, '11.1  Prasyarat Backend Odoo', level=2)
    prereqs = [
        'Modul grt_sales_business_category sudah ter-install dan aktif di Odoo.',
        'Modul grt_inventory_business_category aktif (untuk fitur produk per business category).',
        'CORS dikonfigurasi di Nginx/Apache untuk mengizinkan request dari domain frontend.',
        'Cookie SameSite dan Secure dikonfigurasi sesuai dengan skema domain (HTTPS untuk production).',
        'User Odoo yang digunakan memiliki akses ke menu Sales dan model yang diperlukan.',
    ]
    for p_text in prereqs:
        p = doc.add_paragraph(p_text, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    add_heading(doc, '11.2  Setup Business Category', level=2)
    doc.add_paragraph(
        'Untuk fitur Susu Olahan dan Minimarket, buat master Business Category di Odoo:'
    )
    bc_steps = [
        'Masuk ke Sales Odoo.',
        'Buka menu Business Categories.',
        'Buat category baru dengan Name = SUSU OLAHAN dan isi Code (contoh: SUSU_OLAHAN).',
        'Pilih Company dan isi Analytic Account jika diperlukan.',
        'Aktifkan "Gunakan Ongkos Kirim di Sales Order" jika order membutuhkan auto ongkir.',
        'Set produk susu kemasan ke Business Category SUSU OLAHAN.',
        'Buat kategori produk All / Saleable / Ongkos Kirim dan set produk ongkir ke sana.',
        'Buka Sales > Frontend Shipping Rules dan buat rule ongkir per wilayah.',
    ]
    for i, step in enumerate(bc_steps, 1):
        p = doc.add_paragraph(f'{i}. {step}')
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    add_heading(doc, '11.3  Menjalankan Aplikasi Frontend', level=2)
    add_code_block(doc,
        '# Install dependencies\n'
        'npm install\n\n'
        '# Development server\n'
        'npm run dev\n\n'
        '# Build untuk production\n'
        'npm run build'
    )

    # ===== 12. CATATAN KEAMANAN =====
    add_heading(doc, '12. Catatan Keamanan', level=1)
    doc.add_paragraph(
        'Aplikasi ini mengikuti praktik keamanan standar untuk aplikasi web yang berinteraksi '
        'dengan backend Odoo:'
    )
    security_notes = [
        'Autentikasi berbasis session Odoo — tidak ada token/password yang disimpan di localStorage. Hanya data session non-sensitif (uid, nama, company) yang di-persist.',
        'Semua request menggunakan credentials: "include" untuk mengirim cookie session secara otomatis.',
        'Tidak ada hardcoded URL atau credentials di dalam source code — base URL diisi oleh pengguna pada form login.',
        'Input QR diekstrak dan divalidasi sebelum dikirim ke API (extractCustomerQrRef).',
        'Error dari API tidak mengekspos detail internal backend ke pengguna akhir.',
        'Navigation guard mencegah akses ke halaman terproteksi tanpa session yang valid.',
        'Untuk production, pastikan HTTPS diaktifkan dan konfigurasi CORS di backend dibatasi hanya untuk domain frontend yang diizinkan.',
    ]
    for note in security_notes:
        p = doc.add_paragraph(note, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)

    # ===== PENUTUP =====
    doc.add_paragraph()
    add_info_box(doc,
        'Dokumen ini dibuat secara otomatis berdasarkan source code aplikasi vue-odoo-sales. '
        f'Versi dokumen: 1.0  |  Tanggal: {datetime.date.today().strftime("%d %B %Y")}  |  '
        'Modul backend: grt_sales_business_category (Odoo)',
        bg_color='E3F2FD'
    )

    # Save
    output_path = r'c:\projek\vue-odoo-sales\Dokumentasi_Fungsional_Vue_Odoo_Sales.docx'
    doc.save(output_path)
    print(f'Dokumen berhasil dibuat: {output_path}')

if __name__ == '__main__':
    build_document()
